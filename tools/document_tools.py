import os
import asyncio
import re
from typing import Dict, Any, Optional, List
from pathlib import Path
import markdown
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
import logging
from datetime import datetime

# Template and localization support
from templates import TemplateManager, create_template_manager
from localization import LanguageManager, create_language_manager

logger = logging.getLogger(__name__)


class ContentDeduplicator:
    """Tool for removing duplicate headings and content from markdown documents"""

    def __init__(self):
        self.heading_pattern = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)

    def remove_duplicate_headings(self, content: str) -> str:
        """Remove duplicate consecutive headings in markdown content"""
        if not content:
            return content

        lines = content.split('\n')
        processed_lines = []
        previous_headings = {}  # level -> text mapping

        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # Check if current line is a heading
            heading_match = self.heading_pattern.match(line)
            if heading_match:
                level = len(heading_match.group(1))  # Number of # symbols
                text = heading_match.group(2).strip()

                # Check if this is a duplicate of the previous heading at the same or higher level
                should_skip = False

                # Check for exact duplicate at same level
                if level in previous_headings and previous_headings[level] == text:
                    should_skip = True
                    logger.info(f"Removing duplicate heading: {line}")

                # Check for duplicate at different levels (e.g., ## Title followed by ### Title)
                for prev_level, prev_text in previous_headings.items():
                    if prev_text == text and abs(prev_level - level) == 1:
                        # Keep the higher level heading (fewer #)
                        if level > prev_level:
                            should_skip = True
                            logger.info(f"Removing lower-level duplicate heading: {line}")
                        break

                if not should_skip:
                    processed_lines.append(lines[i])
                    # Update previous headings, clearing lower levels
                    previous_headings = {k: v for k, v in previous_headings.items() if k < level}
                    previous_headings[level] = text
                else:
                    # Skip this line and any immediately following empty lines
                    while i + 1 < len(lines) and not lines[i + 1].strip():
                        i += 1
            else:
                processed_lines.append(lines[i])
                # Reset heading tracking if we encounter non-heading content
                if line:  # Only reset on non-empty lines
                    previous_headings.clear()

            i += 1

        return '\n'.join(processed_lines)

    def remove_redundant_sections(self, content: str) -> str:
        """Remove redundant sections that appear multiple times"""
        if not content:
            return content

        # Split content into sections based on headings
        sections = re.split(r'\n(?=#{1,6}\s)', content)
        unique_sections = []
        seen_section_titles = set()

        for section in sections:
            if not section.strip():
                continue

            # Extract section title
            lines = section.split('\n')
            first_line = lines[0].strip()
            heading_match = self.heading_pattern.match(first_line)

            if heading_match:
                title = heading_match.group(2).strip().lower()

                # Check for similar titles (fuzzy matching)
                is_duplicate = False
                for seen_title in seen_section_titles:
                    if self._are_titles_similar(title, seen_title):
                        is_duplicate = True
                        logger.info(f"Removing duplicate section: {first_line}")
                        break

                if not is_duplicate:
                    unique_sections.append(section)
                    seen_section_titles.add(title)
            else:
                # Non-heading section, keep it
                unique_sections.append(section)

        return '\n'.join(unique_sections)

    def _are_titles_similar(self, title1: str, title2: str, threshold: float = 0.8) -> bool:
        """Check if two titles are similar using simple string similarity"""
        if title1 == title2:
            return True

        # Simple similarity check based on common words
        words1 = set(title1.split())
        words2 = set(title2.split())

        if not words1 or not words2:
            return False

        intersection = words1.intersection(words2)
        union = words1.union(words2)

        similarity = len(intersection) / len(union)
        return similarity >= threshold

    def clean_content(self, content: str) -> str:
        """Main method to clean content by removing duplicates"""
        if not content:
            return content

        # First remove duplicate headings
        content = self.remove_duplicate_headings(content)

        # Then remove redundant sections
        content = self.remove_redundant_sections(content)

        # Clean up extra whitespace
        content = re.sub(r'\n{3,}', '\n\n', content)

        return content.strip()


class DocumentProcessorTool:
    def __init__(self, output_dir: str = "./outputs"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.deduplicator = ContentDeduplicator()

    def process_markdown(self, content: str) -> str:
        # Convert markdown to HTML for better processing
        html_content = markdown.markdown(content, extensions=['tables', 'toc', 'codehilite'])
        return html_content

    def clean_content_for_export(self, content: str) -> str:
        # First remove duplicate headings and sections
        content = self.deduplicator.clean_content(content)

        # Then remove or replace problematic characters for document export
        content = content.replace('\u2019', "'")  # Replace smart quotes
        content = content.replace('\u201c', '"')
        content = content.replace('\u201d', '"')
        content = content.replace('\u2013', '-')  # En dash
        content = content.replace('\u2014', '--')  # Em dash
        return content
    
    def generate_filename(self, title: str, format_type: str, timestamp: bool = True) -> str:
        # Clean title for filename
        clean_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        clean_title = clean_title.replace(' ', '_')
        
        if timestamp:
            timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{clean_title}_{timestamp_str}.{format_type}"
        else:
            filename = f"{clean_title}.{format_type}"
        
        return filename


class PDFGeneratorTool(DocumentProcessorTool):
    def __init__(self, output_dir: str = "./outputs"):
        super().__init__(output_dir)
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        # Create custom styles
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Title'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center alignment
        ))
        
        self.styles.add(ParagraphStyle(
            name='CustomHeading1',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceAfter=12,
            spaceBefore=12
        ))
        
        self.styles.add(ParagraphStyle(
            name='CustomHeading2',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceAfter=10,
            spaceBefore=10
        ))
    
    async def generate_pdf(self, content: str, title: str, metadata: Optional[Dict[str, Any]] = None) -> str:
        try:
            filename = self.generate_filename(title, "pdf")
            filepath = self.output_dir / filename
            
            # Clean content
            content = self.clean_content_for_export(content)
            
            # Create PDF document
            doc = SimpleDocTemplate(
                str(filepath),
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Build content
            story = []
            
            # Add title
            story.append(Paragraph(title, self.styles['CustomTitle']))
            story.append(Spacer(1, 20))
            
            # Add metadata if provided
            if metadata:
                date = metadata.get('date', datetime.now().strftime("%Y-%m-%d"))
                date_label = metadata.get('date_label', 'Generated on')
                story.append(Paragraph(f"{date_label}: {date}", self.styles['Normal']))
                story.append(Spacer(1, 20))
            
            # Process content by sections
            sections = content.split('\n\n')
            for section in sections:
                if section.strip():
                    # Check if it's a heading
                    if section.startswith('# '):
                        heading_text = section[2:].strip()
                        story.append(Paragraph(heading_text, self.styles['CustomHeading1']))
                    elif section.startswith('## '):
                        heading_text = section[3:].strip()
                        story.append(Paragraph(heading_text, self.styles['CustomHeading2']))
                    else:
                        # Regular paragraph
                        story.append(Paragraph(section.strip(), self.styles['Normal']))
                    
                    story.append(Spacer(1, 12))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"PDF generated successfully: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Failed to generate PDF: {str(e)}")
            raise


class DocxGeneratorTool(DocumentProcessorTool):
    def __init__(self, output_dir: str = "./outputs"):
        super().__init__(output_dir)
    
    async def generate_docx(self, content: str, title: str, metadata: Optional[Dict[str, Any]] = None) -> str:
        try:
            filename = self.generate_filename(title, "docx")
            filepath = self.output_dir / filename
            
            # Clean content
            content = self.clean_content_for_export(content)
            
            # Create document
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = 1  # Center alignment
            
            # Add metadata if provided
            if metadata:
                date = metadata.get('date', datetime.now().strftime("%Y-%m-%d"))
                date_label = metadata.get('date_label', 'Generated on')
                doc.add_paragraph(f"{date_label}: {date}")
                doc.add_paragraph("")  # Empty line
            
            # Process content by sections
            sections = content.split('\n\n')
            for section in sections:
                if section.strip():
                    # Check if it's a heading
                    if section.startswith('# '):
                        heading_text = section[2:].strip()
                        doc.add_heading(heading_text, level=1)
                    elif section.startswith('## '):
                        heading_text = section[3:].strip()
                        doc.add_heading(heading_text, level=2)
                    elif section.startswith('### '):
                        heading_text = section[4:].strip()
                        doc.add_heading(heading_text, level=3)
                    else:
                        # Regular paragraph
                        doc.add_paragraph(section.strip())
            
            # Save document
            doc.save(str(filepath))
            
            logger.info(f"DOCX generated successfully: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Failed to generate DOCX: {str(e)}")
            raise


class MarkdownGeneratorTool(DocumentProcessorTool):
    def __init__(self, output_dir: str = "./outputs"):
        super().__init__(output_dir)
    
    async def generate_markdown(self, content: str, title: str, metadata: Optional[Dict[str, Any]] = None) -> str:
        try:
            filename = self.generate_filename(title, "md")
            filepath = self.output_dir / filename

            # Build markdown content
            markdown_content = []

            # Add title
            markdown_content.append(f"# {title}\n")

            # Add metadata if provided
            if metadata:
                date = metadata.get('date', datetime.now().strftime("%Y-%m-%d"))
                date_label = metadata.get('date_label', 'Generated on')
                sources_label = metadata.get('sources_label', 'Sources')

                markdown_content.append(f"**{date_label}:** {date}\n")

                if metadata.get('sources'):
                    markdown_content.append(f"## {sources_label}\n")
                    for i, source in enumerate(metadata['sources'], 1):
                        markdown_content.append(f"{i}. {source}\n")

                markdown_content.append("---\n")
            
            # Add main content
            markdown_content.append(content)
            
            # Write to file
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write('\n'.join(markdown_content))
            
            logger.info(f"Markdown generated successfully: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Failed to generate Markdown: {str(e)}")
            raise


class DocumentPublisher:
    def __init__(self, output_dir: str = "./outputs",
                 template_manager: Optional[TemplateManager] = None,
                 language_manager: Optional[LanguageManager] = None):
        self.output_dir = output_dir
        self.pdf_generator = PDFGeneratorTool(output_dir)
        self.docx_generator = DocxGeneratorTool(output_dir)
        self.markdown_generator = MarkdownGeneratorTool(output_dir)

        self.template_manager = template_manager or create_template_manager()
        self.language_manager = language_manager or create_language_manager()

    async def publish_all_formats(self, content: str, title: str,
                                 formats: Dict[str, bool],
                                 metadata: Optional[Dict[str, Any]] = None,
                                 template_name: str = "default",
                                 language_code: str = "en") -> Dict[str, str]:
        results = {}
        tasks = []

        # Apply template and localization
        processed_content, processed_title, processed_metadata = await self._apply_template_and_localization(
            content, title, metadata, template_name, language_code
        )

        if formats.get("markdown", False):
            tasks.append(("markdown", self.markdown_generator.generate_markdown(
                processed_content, processed_title, processed_metadata)))

        if formats.get("pdf", False):
            tasks.append(("pdf", self.pdf_generator.generate_pdf(
                processed_content, processed_title, processed_metadata)))

        if formats.get("docx", False):
            tasks.append(("docx", self.docx_generator.generate_docx(
                processed_content, processed_title, processed_metadata)))

        # Execute all tasks concurrently
        for format_name, task in tasks:
            try:
                filepath = await task
                results[format_name] = filepath
            except Exception as e:
                logger.error(f"Failed to generate {format_name}: {str(e)}")
                results[format_name] = f"Error: {str(e)}"

        return results

    async def _apply_template_and_localization(self, content: str, title: str,
                                             metadata: Optional[Dict[str, Any]],
                                             template_name: str,
                                             language_code: str) -> tuple[str, str, Dict[str, Any]]:
        try:
            # Set language first
            self.language_manager.set_language(language_code)
            lang_config = self.language_manager.get_current_language()

            # Check if template should be used
            template = None
            if self.template_manager.should_use_template(template_name):
                template = self.template_manager.get_template(template_name)
                if not template:
                    template = self.template_manager.get_default_template()

            # Process title
            processed_title = title

            # Process metadata with localization
            processed_metadata = metadata.copy() if metadata else {}
            if lang_config:
                # Localize metadata labels
                processed_metadata["date_label"] = self.language_manager.translate("generated_on")
                processed_metadata["sources_label"] = self.language_manager.translate("sources")
                processed_metadata["date_format"] = lang_config.date_format

                # Format date according to language
                if "date" in processed_metadata:
                    try:
                        date_obj = datetime.fromisoformat(processed_metadata["date"])
                        processed_metadata["date"] = date_obj.strftime(lang_config.date_format)
                    except:
                        pass

            # Process content with template structure (only if template is specified)
            if template:
                processed_content = await self._apply_template_structure(content, template, lang_config)
            else:
                # No template - apply only basic localization to existing content
                processed_content = await self._apply_basic_localization(content, lang_config)

            return processed_content, processed_title, processed_metadata

        except Exception as e:
            logger.warning(f"Failed to apply template and localization: {str(e)}")
            return content, title, metadata or {}

    async def _apply_template_structure(self, content: str, template, lang_config) -> str:
        try:
            if not template or not template.sections:
                return content

            # Split content into sections (basic implementation)
            content_sections = content.split('\n## ')
            if len(content_sections) > 1:
                content_sections[0] = content_sections[0].replace('# ', '')
                for i in range(1, len(content_sections)):
                    content_sections[i] = '## ' + content_sections[i]

            # Apply template section titles
            processed_sections = []
            for i, section_template in enumerate(sorted(template.sections, key=lambda x: x.order)):
                if i < len(content_sections):
                    section_content = content_sections[i]

                    # Get localized section title
                    localized_title = self.language_manager.get_section_title(
                        section_template.name, lang_config.code if lang_config else "en"
                    )

                    # Replace section header with localized title
                    if section_content.startswith('## '):
                        section_content = f"## {localized_title}\n" + '\n'.join(section_content.split('\n')[1:])
                    elif i == 0:
                        section_content = f"## {localized_title}\n{section_content}"

                    processed_sections.append(section_content)

            # Add any remaining content sections
            if len(content_sections) > len(template.sections):
                processed_sections.extend(content_sections[len(template.sections):])

            return '\n\n'.join(processed_sections)

        except Exception as e:
            logger.warning(f"Failed to apply template structure: {str(e)}")
            return content

    async def _apply_basic_localization(self, content: str, lang_config) -> str:
        """Apply basic localization without template structure"""
        try:
            if not lang_config:
                return content

            # Check if content is already localized by looking for common localized patterns
            # If content already contains localized titles, skip further localization
            localized_patterns = [
                self.language_manager.translate('introduction'),
                self.language_manager.translate('conclusion'),
                self.language_manager.translate('references')
            ]

            # If any localized pattern is found, assume content is already localized
            for pattern in localized_patterns:
                if pattern in content and pattern != pattern.replace("_", " ").title():
                    logger.info("Content appears to be already localized, skipping basic localization")
                    return content

            # Basic section title translation for common patterns
            common_sections = {
                "# Introduction": f"# {self.language_manager.translate('introduction')}",
                "# Methodology": f"# {self.language_manager.translate('methodology')}",
                "# Results": f"# {self.language_manager.translate('results')}",
                "# Discussion": f"# {self.language_manager.translate('discussion')}",
                "# Conclusion": f"# {self.language_manager.translate('conclusion')}",
                "# References": f"# {self.language_manager.translate('references')}",
                "## Introduction": f"## {self.language_manager.translate('introduction')}",
                "## Methodology": f"## {self.language_manager.translate('methodology')}",
                "## Results": f"## {self.language_manager.translate('results')}",
                "## Discussion": f"## {self.language_manager.translate('discussion')}",
                "## Conclusion": f"## {self.language_manager.translate('conclusion')}",
                "## References": f"## {self.language_manager.translate('references')}"
            }

            processed_content = content
            for english_title, localized_title in common_sections.items():
                processed_content = processed_content.replace(english_title, localized_title)

            return processed_content

        except Exception as e:
            logger.warning(f"Failed to apply basic localization: {str(e)}")
            return content
